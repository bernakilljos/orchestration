"""팀장 보고 docx 빌더 — 2026-08-21."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUT = Path(__file__).with_name("2026-08-21-work-report-to-manager.docx")

doc = Document()

style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(10.5)

sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2)
sec.right_margin = Cm(2)


def h1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def h2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def h3(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11.5)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def hr():
    p = doc.add_paragraph()
    p.add_run("─" * 60).font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


def table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r in rows:
        cells = t.add_row().cells
        for i, val in enumerate(r):
            cells[i].text = str(val)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    doc.add_paragraph()


# ===== 헤더 =====
h1("업무 현황 및 계획 보고")
info = doc.add_paragraph()
info.add_run("보고자: ").bold = True
info.add_run("서성종 (아이티센코어 ESG사업부 개발팀 PL)\n")
info.add_run("작성일: ").bold = True
info.add_run("2026-08-21\n")
info.add_run("수신: ").bold = True
info.add_run("팀장님 (내일까지 회신 요청 건)")
hr()

# ===== 1. 현재 =====
h2("1. 현재 하고 있는 일")

h3("사업부 사업 영역 (3 축)")
table(
    ["축", "규제·법령", "주력 제품", "우리 개발팀 담당"],
    [
        ["① 내부회계 (K-SOX)", "신외감법 · 내부회계관리제도", "MicroICM · MicroICM@Cloud · ICM Agent (Go)", "개발 리드"],
        ["② 리스크모니터링 (RMS)", "자금·구매·법인카드 상시감사", "MicroRisk-X (신사업) · Cen's TRM · MicroEWS", "PL (v14 저자)"],
        ["③ CCP (준법경영)", "준법지원인·컴플라이언스 프로그램", "(기존 자산 활용 + 신규 검토)", "개발 지원"],
    ],
    col_widths=[3.5, 4, 6, 3.5],
)
p_rms = doc.add_paragraph()
p_rms.add_run("RMS 상세 (담당 파트): ").bold = True
p_rms.add_run("외부 I/F 로 법인카드·자금·구매 등 거래 데이터 수집 → 이상 거래·리스크 지표·규정 위반 자동 감지·실시간 모니터링.")

h3("병행 프로젝트 (4 트랙)")
table(
    ["#", "프로젝트", "위치", "성격", "역할"],
    [
        ["1", "연결결산 (Contabulo)", "C:\\IFRS_PJT", "Node.js 웹앱 (Server + Client)", "개발 리드"],
        ["2", "RMS (MicroRisk-X 신사업)", "C:\\RMS_PJT", "GRC 백엔드(grc-ba) + UI(grc-ui) · PPT v14 저자", "PL"],
        ["3", "ICM Agent (Go)", "C:\\ICMAI_PJT\\Project", "MicroICM 확장 · AI 에이전트 · Go 언어", "개발 리드"],
        ["4", "AI 오케스트레이션 킷 v1", "C:\\pjt\\orchestration_v1", "멀티AI (Claude+Codex+Gemini) 사내 자동화 인프라", "저자·유지보수"],
    ],
    col_widths=[0.8, 4.5, 3.5, 5.5, 2.2],
)

h3("이번 달 완료 (2026-08 실적)")
bullet("MicroRisk-X 신사업 PPT v14 초안")
bullet("Orchestration Kit v1 신기술 반영 (Opus 5 · Sonnet 5 · Managed Agents · Qwen 3.8-27B)")
bullet("사내 조사 규율 룰 6종 신설 (환경 의존 결함 · 계측 3축 · 조사 방법론)")
bullet("산업 ML RMS 이식 종합지도 pptx 갱신")

h3("도메인 강점 활용")
bullet("ISMS-P 자격 3년 (개인 도메인)")
bullet("ESG (Governance) 안 정보보호공시 자연 연결")
bullet("아이티센코어 ESG사업부 6대 자원 (MicroICM · Cen's TRM · MicroEWS · MicroBSC · MicroICM@Cloud · Contabulo) 연동")

hr()

# ===== 2. 어려운 점 =====
h2("2. 어려운 점")

h3("근본 병목 (개발팀 관점)")
bullet("개인 병목 — 사업부 3축 (K-SOX·RMS·CCP) 을 개발팀이 모두 커버. 4 트랙 병행 · PL 부재 시 진행 정지 리스크")
bullet("개발팀이 자산을 만들어야 매출이 나오는 구조 — 영업·컨설팅이 팔 근거를 시기별로 릴리스해야 하는데 시간 부족")
bullet("하루 각 프로젝트 2~3시간이 실질 한계")

h3("기술 부담")
bullet("2026 하반 AI 파도 — Claude Opus 5 · Sonnet 5 · Fable 5 · Qwen 3.8-27B · Managed Agents · Inference hooks 등 6주 새 대형 변화 8건+")
bullet("ICM Agent Go 새 언어 병행 (기존 Python/Node 스택과 다름)")
bullet("개인정보·기업 민감정보 처리 리스크 — RMS I/F 데이터 (법인카드·자금·구매) 외부 API 사용 = 규제 리스크 → 로컬 AI 필수인데 사내 GPU 미보유")

h3("도메인 공백")
bullet("회계 정합성 — 회계법인 영역 · 팀 자체 커버 X → 외주 필요")
bullet("ISMS-P — PL 개인 자격 3년만 · 팀 강점 X → 컨설팅 파트너십 필요")
bullet("CCP (준법경영) — 신규 진입 영역 · 도메인 지식·시장 조사 시간 필요")

h3("인프라 반복")
bullet("개발·검증 세션 동시 병행 시 조용한 유실 (git worktree 물리 분리 필요)")
bullet("Windows Git Credential Manager GUI 인증 CLI 세션 hang 반복")
bullet("claude.ai 이미지 첨부 배경 탭 제약 (2026-08-20 postmortem 완료)")
bullet("사내 개발 표준 환경·CI/CD 부재로 반복 인프라 문제 발생")

hr()

# ===== 3. 건의사항 =====
h2("3. 건의사항")
p_intro3 = doc.add_paragraph()
p_intro3.add_run("원칙: ").bold = True
p_intro3.add_run("개발팀이 자산을 만들어놓아야 실적으로 이어짐. 이를 위한 지원 요청.")

h3("인력 (실적 리스크 완화)")
bullet("① RMS 또는 IFRS 백업 개발자 1명 즉시 배치 — 4 트랙 중 최소 1개 부담 완화 · PL 부재 시 진행 정지 리스크 제거")
bullet("② CCP (준법경영) 도메인 리소스 — 신규 진입 영역 · 도메인 스터디·시장 조사·MVP 프로토타입 시간 확보")
bullet("③ 개발팀 확대 로드맵 — 2027년까지 3~4명 · 프로젝트별 PL 분리 · 신입 온보딩 표준화 (Orchestration Kit 이미 갖춰짐)")

h3("인프라 (원가·수익성 직결)")
bullet("④ 사내 GPU 서버 (Qwen 3.8-27B 로컬 배포용) — RMS I/F 데이터 (법인카드·자금·구매) 외부 API 미노출 대량 판정 가능 · 개인정보·기업 민감정보 규제 대응 + AI API 원가 절감 이중")
bullet("⑤ Orchestration Kit 사내 표준화 지원 — kit 을 다른 부서·계열사 install 배포 시 인프라·라이센스 지원 (Zero-touch 자동화 실증 사례 축적)")
bullet("⑥ Anthropic Enterprise 계약 검토 — Managed Agents + Inference hooks 로 사내 AI 거버넌스 표준화")

h3("파트너십 (도메인 보강)")
bullet("⑦ ISMS 컨설팅 파트너십 — KISA 인증 컨설팅사 1~2곳 협약 (재판매·리드젠)")
bullet("⑧ 회계법인 외주 관계 — 회계 정합성 영역 보강 · MicroICM/Contabulo 도입 진단 파트너")

h3("방향성 결정")
bullet("⑨ 신사업 (RMS · CCP 신규) vs 유지보수 (K-SOX · IFRS) 리소스 배분 명시 — 팀장님 판단 요청. 개발팀이 무엇을 먼저 완비할지 결정 필요")

hr()

# ===== 4. 향후 1~2년 계획 =====
h2("4. 향후 1~2년 계획 — 개발팀이 실적으로 이어지게 할 수 있는 것 (3축 관점)")

p_intro4 = doc.add_paragraph()
p_intro4.add_run("관점: ").bold = True
p_intro4.add_run("개발팀. 영업·컨설팅·경영이 팔 수 있으려면 우리가 먼저 만들어놓아야 실적으로 이어짐. 사업부 3축 (K-SOX 내부회계 · RMS 리스크모니터링 · CCP 준법경영) 자산을 시기별로 완비.")

h3("4.1 시기별 릴리스 (3축 관점)")
table(
    ["시점", "축", "릴리스", "산출", "실적 연결"],
    [
        ["2026 H2", "RMS", "MicroRisk-X v15", "2026-08 신기술 반영 · I/F 시나리오 3종 (법인카드·자금·구매)", "유료 PoC 3~5곳 (신사업 ARR 씨앗)"],
        ["2026 H2", "K-SOX", "ICM Agent Go GA", "MicroICM 연동 완료 · 성능 데이터 · 시나리오", "MicroICM 고객 프리미엄 옵션 (ARPU 상승)"],
        ["2026 H2", "K-SOX", "Contabulo 클라우드", "IFRS 연결결산 자동화 · MicroICM@Cloud 연동", "상장사·중견 클라우드 신규 · SaaS 전환"],
        ["2026 H2", "CCP", "CCP 준법 자동화 프로토타입", "준법감시 규정 매핑 · CCP 지표 자동 수집 파이프라인", "신규 시장 진입 근거 (3축 완비)"],
        ["2026 H2", "인프라", "Orchestration Kit v2", "Zero-touch 100% · 사내 다른 부서 install 실증", "컨설팅 신뢰 근거 + B2B 검토 근거"],
        ["2027 상반", "RMS", "MicroRisk-X GA", "유상 PoC 결과 정식판 · 파트너 install kit", "유료 고객 5+ ARR · 재판매 파트너"],
        ["2027 상반", "컨설팅", "컨설팅 도구 세트", "install kit · 진단 툴 · 교육 자료", "AI 자동화·ESG 공시·CCP 진단 판매 시작"],
        ["2027 하반", "인프라", "Orchestration Kit B2B", "파일럿 유상 판매 준비 · 라이센스·SLA", "신규 제품 라인"],
        ["2027 하반", "규제", "규제 대응 자동화 세트", "정보보호공시·ESG 공시·K-SOX 강화·CCP 준법 자동화", "규제 대응 SaaS 신규 (자산 100억 이상 확대)"],
    ],
    col_widths=[1.5, 1.5, 3.5, 5.5, 5],
)

h3("4.2 개발팀이 만들어야 매출로 이어지는 5가지 (팔 수 있는 근거)")
para("우리가 만들지 않으면 아무도 팔 수 없는 것:")
bullet("① 사내 자동화 100% 실증 사례 — Orchestration Kit 을 우리 사업부·다른 부서에 배포한 실측. '우리는 이미 해봤다' 가 컨설팅·B2B 유일한 신뢰 근거")
bullet("② 3축 제품 데모 시나리오 각 3종 — K-SOX (MicroICM·Contabulo·ICM Agent) · RMS (MicroRisk-X · I/F 법인카드·자금·구매) · CCP (준법감시) 별 규제 대응 시나리오. 영업 즉시 시연")
bullet("③ 3축 크로스셀 기술 근거 — MicroICM 고객 → MicroRisk-X → CCP 데이터 연동·SSO·지표 매핑을 미리 준비. 3축 통합 판매 근거는 개발이 만든다")
bullet("④ 파트너 install kit — ISMS·회계·준법 파트너 자체 판매용 install·training. 파트너 매출은 우리가 준비한 만큼 나옴")
bullet("⑤ 규제 대응 자동화 세트 — 정보보호공시·ESG 공시·K-SOX 강화·준법경영 규제 각각 대응. 규제 시점 = 매출 기회. 개발 준비 없으면 놓침")

h3("4.3 개발팀이 원가·수익성에 직접 기여")
bullet("API 원가 절감 + 규제 대응 이중 효과: Qwen 3.8-27B 로컬 (Apache 2.0 · 262K context · SWE-bench Pro 61.7%). RMS I/F 데이터 (법인카드·자금·구매) 는 개인정보·기업 민감정보 → 외부 API 미노출 필수. 개발팀 사내 GPU 셋업 = 원가 + 규제 이중")
bullet("반복 SI 자동화: Orchestration Kit 흡수로 인당 처리량 상승 · 매출 대비 원가율 개선 · 개발팀 매출 기여도가 숫자로 나옴")
bullet("3축 KPI 대시보드: 프로젝트별·3축별 매출·수익성 대시보드 (사내 도구). 경영 판단 근거를 우리가 제공")

h3("4.4 그룹 채널 (개발이 install · 영업이 성사)")
bullet("역할 분담: 아이티센 그룹 계열사간 거래는 영업 성사 · 개발팀은 install·튜닝·유지보수 리소스 준비")
bullet("성과 이중: 그룹 내부 매출 안정 (즉시) + 외부 판매 레퍼런스 (중장기)")

h3("4.5 2028 시장 선점 기술 준비 (지금 시작해야 그때 팜)")
bullet("정보보호공시 자동화 SaaS — 자산 100억 이상 확대 규제 · 지금 프로토타입 → 2027 하반 SaaS")
bullet("Post-Quantum Security (양자내성보안 · CRYSTALS-Kyber·Dilithium·SPHINCS+) — 금융·공공 진입 근거 · 도입 진단 툴 · NIST 표준 반영")
bullet("Synthetic Data (합성데이터 · Gretel·Mostly AI·SDV) — RMS I/F · CCP 학습용 · 개인정보 없는 데이터 판매 · 규제 대응 이중")

h3("§ 4 요약 (개발팀 관점)")
p_sum4 = doc.add_paragraph()
p_sum4.add_run("개발이 만들어 놓아야 영업·경영이 팔 수 있다. 사업부 3축 (K-SOX · RMS · CCP) 자산을 시기별로 완비.").bold = True
table(
    ["개발팀이 만들 것", "그것이 이어질 실적"],
    [
        ["시기별 릴리스 9건 (4.1) · 3축 모두", "K-SOX ARPU · RMS ARR · CCP 진입 · SaaS 전환 · B2B 도구"],
        ["매출 재료 5가지 (4.2)", "컨설팅 신뢰 · 3축 시연 · 크로스셀 성사 · 파트너 매출 · 규제 대응"],
        ["원가·수익 기여 3건 (4.3)", "API 원가 절감 · 처리량 상승 · 경영 3축 KPI 근거"],
        ["그룹 install 리소스 (4.4)", "그룹 내부 매출 + 외부 레퍼런스"],
        ["2028 기술 준비 3영역 (4.5)", "시장 선점 · 규제 대응 신규 상품"],
    ],
    col_widths=[6, 11],
)
p_msg = doc.add_paragraph()
p_msg.add_run("핵심 메시지: ").bold = True
p_msg.add_run("개발팀은 영업·경영·컨설팅을 직접 하지 않지만, 팔 수 있는 자산을 시기별로 만들어놓는 것이 개발팀 실적의 본질. 위 5개가 준비되면 사업부 3축이 영업·경영·컨설팅으로 각자 매출 확대.")

hr()

# ===== 5. 요약 =====
h2("5. 요약")
bullet("사업부 3축 (K-SOX · RMS · CCP) 을 개발팀 4 트랙으로 커버 · PL 개인 병목이 최대 리스크")
bullet("AI 자동화 인프라 (Orchestration Kit v1) 는 사내 자산 · 부서 배포 → 컨설팅·B2B 상품화 가능")
bullet("RMS I/F 데이터 (법인카드·자금·구매) = 사내 GPU (Qwen 로컬) 필수 · 규제 대응 + API 원가 절감 이중")
bullet("ESG G(overnance) + AI 자동화 조합 = 사업부 진입 근거 · 2027 상용화 목표")
bullet("회계·ISMS·준법 도메인 공백은 외부 파트너십으로 보강이 정직한 접근")
p_ask = doc.add_paragraph()
p_ask.add_run("팀장님께 즉시 요청 3가지: ").bold = True
p_ask.add_run("① 백업 개발자 · ② 사내 GPU 서버 · ③ 신사업 vs 유지보수 우선순위 결정")

hr()

footer = doc.add_paragraph()
footer.add_run("보고자: ").bold = True
footer.add_run("서성종 · sjseo@itcen.com · 아이티센코어 ESG사업부 개발팀")

doc.save(OUT)
print(f"OK · {OUT}")
